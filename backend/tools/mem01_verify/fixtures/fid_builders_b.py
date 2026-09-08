"""
Role: In-memory builders for the FID synthetic ORIGINALS — part 2: the OOXML containers. Renders
      a `DocSpec` into a real .docx (ordered paragraphs, list items, a real `w:tbl` table, real
      external `w:hyperlink` relationships, an optional page break) and a real .xlsx workbook
      (an opening narrative sheet, one sheet per table carrying a caption row, a header row and
      string data cells, and a closing sheet with real cell hyperlink relationships). Both go
      through `deterministic_zip`, so equal arguments always produce equal payload bytes.
Used by: `fid_cases_c.py` (the docx cases), `fid_cases_d.py` (the xlsx cases and the docx
         embedded in a TNEF container), `fid_cases_f.py` (encoding cases authored in docx/xlsx).
         `deterministic_zip` lives here because only the OOXML containers need it.
Depends on: `python-docx` and `openpyxl` (both declared runtime dependencies of the backend),
            stdlib `zipfile`, and the sibling `fid_cases_a` for the `DocSpec` vocabulary. Nothing
            from `backend/app/`, nothing from `backend/tests/`.
Key invariants:
  - The document body is written in the SPEC's frozen order (paragraphs, list items, table
    caption + `w:tbl`, link paragraphs, trailing paragraphs) — the same order
    `expectation_for_doc` derives its order constraints from.
  - Byte-determinism: document properties are pinned to `FIXED_TIMESTAMP` and the container is
    repacked through `deterministic_zip`, so two calls with equal arguments produce equal bytes.
  - A link is a REAL external relationship, not label text with the URL typed next to it: the
    destination lives only in the relationship, exactly as in a real document, so a text handoff
    that cannot carry it fails the link pair honestly.
  - Every xlsx data cell is written as a STRING. A float or date cell's text form is a rendering
    choice the ORIGINAL does not fix, so a required unit is never a typed numeric cell.
"""

from __future__ import annotations

import zipfile
from collections.abc import Mapping, Sequence
from datetime import UTC, datetime
from io import BytesIO

import docx
import openpyxl
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from tools.mem01_verify.fixtures.fid_cases_a import DocSpec

#: Fixed instant stamped into every container's zip member headers and document properties.
FIXED_TIMESTAMP = (1980, 1, 1, 0, 0, 0)

#: The same instant as a naive datetime (the document-property API takes naive values).
FIXED_DATETIME = datetime(*FIXED_TIMESTAMP, tzinfo=UTC).replace(tzinfo=None)

#: Author stamped into every container — synthetic, never a real person.
FIXTURE_AUTHOR = "MEM-01 fixture generator"

#: The workbook sheets that carry an ORIGINAL's narrative blocks, before and after its tables.
OPENING_SHEET = "Начало"
CLOSING_SHEET = "Приложения"


def _append_external_hyperlink(paragraph: object, label: str, url: str) -> None:
    """Append a real external `w:hyperlink` run carrying `label` to a docx paragraph.

    python-docx has no public hyperlink API, so the relationship and the `w:hyperlink` element are
    written directly; `paragraph.part` and `paragraph._p` are the documented internal seams the
    library's own extension recipes use.
    """
    relationship_id = paragraph.part.relate_to(  # type: ignore[attr-defined]
        url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    text_element = OxmlElement("w:t")
    text_element.text = label
    run.append(text_element)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)  # type: ignore[attr-defined]  # noqa: SLF001


def build_docx(spec: DocSpec, *, page_break_before_tables: bool = False) -> bytes:
    """Render one ORIGINAL as a real .docx in the specification's frozen block order.

    Args:
        spec: the ORIGINAL.
        page_break_before_tables: start a second page before the first table, so the case carries
            a genuine page boundary whose reading order must survive.

    Returns:
        The byte-deterministic .docx container.
    """
    document = docx.Document()
    for text in spec.paragraphs:
        document.add_paragraph(text)
    for bullet in spec.bullets:
        document.add_paragraph(bullet, style="List Bullet")
    if page_break_before_tables and spec.tables:
        document.add_page_break()
    for table_spec in spec.tables:
        document.add_paragraph(table_spec.caption)
        table = document.add_table(rows=1 + len(table_spec.rows), cols=len(table_spec.headers))
        for column_index, header in enumerate(table_spec.headers):
            table.rows[0].cells[column_index].text = header
        for row_index, values in enumerate(table_spec.rows, start=1):
            for column_index, value in enumerate(values):
                table.rows[row_index].cells[column_index].text = value
    for link in spec.links:
        _append_external_hyperlink(document.add_paragraph(), link.label, link.destination)
    for text in spec.trailing:
        document.add_paragraph(text)

    properties = document.core_properties
    properties.author = FIXTURE_AUTHOR
    properties.last_modified_by = FIXTURE_AUTHOR
    properties.created = FIXED_DATETIME
    properties.modified = FIXED_DATETIME
    properties.revision = 1
    buffer = BytesIO()
    document.save(buffer)
    return deterministic_zip(buffer.getvalue())


def build_xlsx(
    sheets: Sequence[tuple[str, Sequence[Sequence[str]]]],
    hyperlinks: Mapping[str, Mapping[str, str]] | None = None,
) -> bytes:
    """Assemble a real .xlsx workbook: one sheet per entry, every value a string cell.

    Args:
        sheets: `(sheet_name, rows)` in workbook order; each row is a sequence of string values
            appended left to right.
        hyperlinks: optional `{sheet_name: {cell_ref: destination}}` — a REAL cell hyperlink
            relationship, the destination living outside the cell's text.

    Returns:
        The byte-deterministic .xlsx container.
    """
    workbook = openpyxl.Workbook()
    workbook.remove(workbook.active)
    for name, rows in sheets:
        worksheet = workbook.create_sheet(title=name)
        for row in rows:
            worksheet.append(list(row))
        for reference, destination in (hyperlinks or {}).get(name, {}).items():
            worksheet[reference].hyperlink = destination
    workbook.properties.creator = FIXTURE_AUTHOR
    workbook.properties.lastModifiedBy = FIXTURE_AUTHOR
    workbook.properties.created = FIXED_DATETIME
    workbook.properties.modified = FIXED_DATETIME
    buffer = BytesIO()
    workbook.save(buffer)
    return deterministic_zip(buffer.getvalue())


def xlsx_from_doc(spec: DocSpec) -> bytes:
    """Render one ORIGINAL as a workbook whose SHEET ORDER mirrors the ORIGINAL's block order.

    Sheet order IS the "page order" unit of the taxonomy for a workbook: an opening narrative
    sheet (paragraphs, list items), then one sheet per table holding its caption row, its header
    row and its data rows, then a closing sheet carrying the link labels (each a REAL cell
    hyperlink whose destination lives outside the cell text) and the trailing paragraphs. That is
    exactly the order `expectation_for_doc` derives its document-flow constraint from.

    Args:
        spec: the ORIGINAL.

    Returns:
        The byte-deterministic .xlsx container.
    """
    sheets: list[tuple[str, Sequence[Sequence[str]]]] = []
    hyperlinks: dict[str, dict[str, str]] = {}
    opening: list[list[str]] = [[text] for text in spec.paragraphs]
    opening.extend([bullet] for bullet in spec.bullets)
    if opening:
        sheets.append((OPENING_SHEET, opening))
    for table in spec.tables:
        rows: list[list[str]] = [[table.caption], list(table.headers)]
        rows.extend(list(row) for row in table.rows)
        sheets.append((table.sheet, rows))
    closing: list[list[str]] = [[link.label] for link in spec.links]
    closing.extend([text] for text in spec.trailing)
    if closing:
        sheets.append((CLOSING_SHEET, closing))
        for offset, link in enumerate(spec.links, start=1):
            hyperlinks.setdefault(CLOSING_SHEET, {})[f"A{offset}"] = link.destination
    if not sheets:
        sheets.append((OPENING_SHEET, [["(empty)"]]))
    return build_xlsx(sheets, hyperlinks)


#: The core-properties member whose timestamps an Office writer stamps from the clock.
_CORE_PROPERTIES_MEMBER = "docProps/core.xml"

#: The document-property timestamps, as they appear in `docProps/core.xml`.
_CORE_TIMESTAMP_TAGS = ("dcterms:created", "dcterms:modified")

#: The W3CDTF spelling of `FIXED_TIMESTAMP` — the instant every container's properties carry.
FIXED_TIMESTAMP_W3CDTF = "1980-01-01T00:00:00Z"


def _normalize_core_properties(member_xml: bytes) -> bytes:
    """Pin the created/modified timestamps of an Office core-properties part to the fixed instant.

    An Office writer overwrites `dcterms:modified` from the clock at save time, so pinning the
    property on the document object is not enough; the value is rewritten here, in the XML.
    """
    text = member_xml.decode("utf-8")
    for tag in _CORE_TIMESTAMP_TAGS:
        while f"<{tag} " in text or f"<{tag}>" in text:
            open_at = text.index(f"<{tag}")
            value_at = text.index(">", open_at) + 1
            close_at = text.index(f"</{tag}>", value_at)
            if text[value_at:close_at] == FIXED_TIMESTAMP_W3CDTF:
                break
            text = text[:value_at] + FIXED_TIMESTAMP_W3CDTF + text[close_at:]
    return text.encode("utf-8")


def deterministic_zip(payload: bytes) -> bytes:
    """Repack a zip container so its bytes depend only on its members' names and contents.

    Office containers (docx, xlsx) stamp the current time into every local file header AND into
    `docProps/core.xml`, which would make one ORIGINAL's payload differ between two calls. This
    rewrites every member with `FIXED_TIMESTAMP`, in sorted name order, at a fixed compression
    level, and pins the core-properties timestamps.

    Args:
        payload: the container bytes as the writing library produced them.

    Returns:
        The byte-deterministic repacked container.
    """
    source_buffer = BytesIO(payload)
    target_buffer = BytesIO()
    with (
        zipfile.ZipFile(source_buffer) as source,
        zipfile.ZipFile(target_buffer, "w", zipfile.ZIP_DEFLATED, compresslevel=6) as target,
    ):
        for name in sorted(source.namelist()):
            info = zipfile.ZipInfo(name, date_time=FIXED_TIMESTAMP)
            info.compress_type = zipfile.ZIP_DEFLATED
            info.external_attr = 0o600 << 16
            data = source.read(name)
            if name == _CORE_PROPERTIES_MEMBER:
                data = _normalize_core_properties(data)
            target.writestr(info, data)
    return target_buffer.getvalue()
