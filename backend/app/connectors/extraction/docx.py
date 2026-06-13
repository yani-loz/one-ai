"""
Role: docx (OOXML WordprocessingML) attachment text extraction (design §2.4) — python-docx
      primary (body paragraphs + tables in DOCUMENT ORDER via Document.iter_inner_content(),
      tables pipe-serialized through the SAME serialize_table as PDF tables), with a stdlib
      zipfile + word/document.xml XML-harvest emergency fallback for OOXML zips python-docx
      rejects. Headers/footers are DELIBERATELY skipped (a §2.4 divergence): per-page running
      headers/footers are boilerplate noise (letterheads, page numbers, confidentiality footers)
      that would repeat into search/embeddings without adding content.
Used by: the IMAP connector's attachment_extractor (dispatches the
         application/vnd.openxmlformats-officedocument.wordprocessingml.document content type);
         scripts.backfill_attachment_extraction (via the seam). The arrow points IN — this module
         imports nothing back from any connector.
Depends on: python-docx (MIT; lxml BSD-3-Clause underneath) — verified GPL-free; sibling extraction
            modules only: app.connectors.extraction.extraction_result (the result contract),
            .common (serialize_table — ONE table serializer for every extractor, never
            re-implemented — MAX_EXTRACTED_CHARS, the shared stored-text cap, + package_version),
            .text_sanitize (sanitize_body_text — the SINGLE stored-text sanitization source, shared
            with email bodies and PDF text); stdlib zipfile + xml.etree (the emergency fallback
            only). Imports NOTHING from any specific connector (connector-agnostic).
Key invariants:
  - extract_docx_text NEVER raises: every failure (zipfile/python-docx/lxml raise diverse
    internals on malformed payloads) degrades to an ExtractionResult; a final catch-all guards
    even our own bugs.
  - Table-cell text appears EXACTLY ONCE: iter_inner_content() yields top-level Paragraph and
    Table blocks in body order — cell paragraphs are reached only through their Table block,
    never doubled by a separate document.paragraphs pass — and MERGED cells (row.cells repeats
    the merge-origin per spanned grid slot) are emitted once, '' for the other slots.
  - PASSWORD-PROTECTED OOXML arrives as an OLE compound file (magic D0 CF 11 E0), not a zip
    (ECMA-376 encryption wraps the package) → `encrypted` with detail
    'ole-wrapped (password-protected)'. No password handling at MVP (design §2.3 policy).
  - Zip-bomb posture — TWO bounds, both BEFORE any parse: (1) the summed DECLARED file_size of
    every zip member is capped at MAX_DECOMPRESSED_BYTES (the raw-BYTES bound: python-docx
    materializes every part's blob at open; the seam's MAX_PARSE_BYTES 50 MB bounds only the
    COMPRESSED payload) → `corrupt` 'zip-expansion-bound'. (2) Members that get DOM-PARSED cost
    far more than their bytes — tiny-element WordprocessingML inflates ~15x its size in lxml
    tree RSS (measured 2026-06-12 on this repo's lxml: 32 MB of empty runs → ~490 MB; at the
    100 MB byte bound that is ~1.5 GB from a <1 MB compressed docx, an OOM SIGKILL the
    never-raise contract cannot catch) — so XML parts get the much tighter MAX_XML_PARTS_BYTES,
    resolved BOTH by member suffix (*.xml / *.rels) AND by [Content_Types].xml content type
    (python-docx picks the part class by CONTENT TYPE, so an Override can disguise an XML part
    behind a media-looking name) → `corrupt` 'xml-expansion-bound'. MAX_DECOMPRESSED_BYTES'
    headroom exists for media members, which are held as raw bytes and never DOM-parsed.
  - No mid-loop char bail (unlike pdf.py's page loop): python-docx parses the WHOLE XML tree at
    Document() open, so a paragraph-loop bail saves nothing — the XML-parts bound is the real
    parse-memory bound, and the post-hoc MAX_EXTRACTED_CHARS cap (→ `truncated`, capped text
    STORED) is the storage bound.
  - `detail` carries exception CLASS NAMES / fixed phrases only — never str(exc) (lxml/zipfile
    messages can embed payload fragments) and never document content. python-docx and lxml were
    LIVE-CHECKED to emit zero log records on parse failure (2026-06-12, this module's leak
    regression test pins it) — no vendor-logger muting needed, unlike pypdf/pdfminer in pdf.py.
  - The emergency XML harvest uses stdlib ElementTree on the size-bounded word/document.xml only;
    expat ≥2.4 (bundled with every supported CPython) has billion-laughs amplification protection
    built in, and ElementTree never resolves external entities.
  - The caller (the seam) enforces the global size ceiling BEFORE this module parses anything.
"""

from __future__ import annotations

import logging
import struct
import zipfile
from io import BytesIO
from xml.etree import ElementTree

from docx import Document
from docx.table import Table

from app.connectors.extraction.common import (
    MAX_EXTRACTED_CHARS,
    serialize_table,
)
from app.connectors.extraction.common import (
    package_version as _package_version,
)
from app.connectors.extraction.extraction_result import (
    STATUS_CORRUPT,
    STATUS_EMPTY,
    STATUS_ENCRYPTED,
    STATUS_EXTRACTED,
    STATUS_TRUNCATED,
    ExtractionResult,
)
from app.connectors.extraction.text_sanitize import sanitize_body_text

logger = logging.getLogger(__name__)

# OLE compound-file magic (D0 CF 11 E0 A1 B1 1A E1): an ECMA-376-encrypted (password-protected)
# OOXML package — or a mislabeled legacy .doc, which is equally unreadable here. Either way the
# honest status is `encrypted` for the declared-docx dispatch (design §2.3/§2.4).
_OLE_MAGIC = b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1"

# The member the stdlib FALLBACK harvests. Deliberately NOT a package gate: OPC resolves the
# main document part via _rels/.rels, and python-docx parses legitimate packages whose main part
# has another name (Word's 'document2.xml' save quirk, non-Word producers) - pre-empting it
# falsely stamped valid documents `corrupt` (2026-06-12 review). A package python-docx rejects
# AND lacking this literal member degrades to `corrupt` in the fallback's own except.
_DOCUMENT_XML = "word/document.xml"

# Central-directory sanity bounds: a zip declaring more members than this is pathological for a
# WordprocessingML package (real docx carry dozens) and only exists to burn CPU/memory in
# directory scans. Enforced from the RAW EOCD record BEFORE ZipFile construction (2026-06-12
# Codex review, measured: a crafted 18 MB / 200k-member zip costs ~12 s CPU + ~115 MB inside
# ZipFile.__init__ alone — the directory parse is the attack, so the gate cannot run after it).
MAX_ZIP_MEMBERS = 4096
# The declared central-directory byte size is bounded too (a liar could declare few entries but
# a huge directory span): 4096 entries x ~0.5 KB of name/extra headroom each.
MAX_CENTRAL_DIR_BYTES = 2 * 1024 * 1024

# End-of-Central-Directory record: signature + the offsets of the fields we read.
# Layout: sig(4) disk(2) cd_disk(2) entries_this_disk(2) TOTAL_ENTRIES(2) CD_SIZE(4) ...
_EOCD_SIGNATURE = b"PK\x05\x06"
_EOCD_MIN_BYTES = 22
_EOCD_MAX_COMMENT = 65535

# Decompressed-expansion ceiling across ALL zip members (declared file_size sum) — the raw-BYTES
# bound. The seam's 50 MB guard bounds compressed bytes only; a crafted bomb compresses far
# beyond 10:1, and python-docx materializes every part's blob at open. The headroom over
# MAX_XML_PARTS_BYTES exists for media members (stored ~1:1, never DOM-parsed).
MAX_DECOMPRESSED_BYTES = 100 * 1024 * 1024

# Declared-size ceiling for the members that get DOM-PARSED (python-docx → lxml; the fallback →
# ElementTree) — the parse-MEMORY bound. Tiny-element WordprocessingML inflates ~15x its byte
# size as an lxml tree (measured 2026-06-12: 32 MB of <w:p><w:r><w:t>a</…> → ~490 MB RSS), so
# 10 MB of XML caps the transient tree near ~150 MB. Legitimate docx XML is low-MB even for
# huge documents (the census's largest docx carry MB-scale document.xml).
MAX_XML_PARTS_BYTES = 10 * 1024 * 1024

# [Content_Types].xml is a KB-scale manifest; the guard parses it (bounded) to resolve which
# members python-docx will DOM-parse. A manifest declaring more than this is itself pathological.
MAX_CONTENT_TYPES_BYTES = 1 * 1024 * 1024

# WordprocessingML namespace, pre-braced for ElementTree tag matching (the fallback harvest).
_WORDML = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"

# OPC content-types manifest: name, pre-braced namespace, and the member suffixes that are
# DOM-parsed regardless of what the manifest claims (the fallback harvest never reads it).
_CONTENT_TYPES_XML = "[Content_Types].xml"
_CONTENT_TYPES_NS = "{http://schemas.openxmlformats.org/package/2006/content-types}"
_XML_MEMBER_SUFFIXES = (".xml", ".rels")

_PYTHON_DOCX = "python-docx"
# The emergency stdlib path has no package version — hand-bump when its rules change so a
# version-aware backfill can target rows harvested under older rules.
_XML_FALLBACK_NAME = "docx-zip-xml"
_XML_FALLBACK_VERSION = "1"


def extract_docx_text(payload: bytes) -> ExtractionResult:
    """Extract a docx's body text; never raises (every failure degrades to a status).

    Pipeline: OLE-magic probe (password-protected OOXML → `encrypted`) → zip validation
    (non-zip/broken-zip → `corrupt`; decompressed-expansion bound → `corrupt`
    'zip-expansion-bound'; member-count bound → `corrupt` 'zip-member-bound';
    DOM-parsed-members bound → `corrupt` 'xml-expansion-bound') → python-docx body traversal in
    document order (paragraphs as lines, tables pipe-serialized via the shared serialize_table;
    headers/footers skipped as boilerplate) → on a python-docx crash, the stdlib
    zipfile + ElementTree harvest of word/document.xml (w:t text, paragraph breaks per w:p;
    design §2.4's emergency path) → both engines failing → `corrupt` with the two exception
    class names in `detail`. Parsed-but-textless → `empty`; text over MAX_EXTRACTED_CHARS →
    `truncated` with the capped text STORED.

    Args:
        payload: the raw docx bytes (the seam has already enforced the global size ceiling).

    Returns:
        An ExtractionResult; text is non-None only for the text-bearing statuses
        (extracted / truncated).
    """
    try:
        return _extract(payload)
    except Exception as unexpected:  # the seam's NEVER-raise contract: even our bugs degrade
        # Class name ONLY - no exc_info: a formatted traceback ends with "Class: str(exc)",
        # and library exception strings can embed payload fragments (2026-06-12 review).
        logger.warning("docx extraction: unexpected failure (%s)", type(unexpected).__name__)
        return ExtractionResult(
            None, STATUS_CORRUPT, detail=f"unexpected:{type(unexpected).__name__}"
        )


def _extract(payload: bytes) -> ExtractionResult:
    """The real pipeline (see extract_docx_text); may raise — the public wrapper degrades."""
    if payload.startswith(_OLE_MAGIC):
        return ExtractionResult(
            None,
            STATUS_ENCRYPTED,
            detail="ole-container (password-protected docx or mislabeled legacy .doc)",
        )
    package_verdict = _validate_ooxml_package(payload)
    if package_verdict is not None:
        return package_verdict
    try:
        body_text = _read_with_python_docx(payload)
    except Exception as docx_error:  # python-docx/lxml raise diverse internals on odd packages
        return _zip_xml_fallback(payload, docx_error)
    text = sanitize_body_text(body_text)
    if not text:
        return ExtractionResult(
            None,
            STATUS_EMPTY,
            extractor_name=_PYTHON_DOCX,
            extractor_version=_package_version(_PYTHON_DOCX),
        )
    return _text_result(text, _PYTHON_DOCX, _package_version(_PYTHON_DOCX))


def _zip_directory_over_bound(payload: bytes) -> bool:
    """EOCD pre-gate: True when the zip DECLARES a pathological central directory.

    Reads the End-of-Central-Directory record straight from the raw tail bytes — total entry
    count (le16 at +10) and central-directory size (le32 at +12) — BEFORE zipfile.ZipFile ever
    parses the directory (construction builds a ZipInfo per entry; the parse itself is the
    resource burn the member bound exists to prevent). The zip64 sentinels (0xFFFF entries /
    0xFFFFFFFF size) mean "real value elsewhere but at least this large", which already exceeds
    both bounds — rejected without parsing the zip64 record. No EOCD found → False: ZipFile
    will raise BadZipFile and the package degrades to `corrupt` through the normal path.
    """
    tail = payload[-(_EOCD_MIN_BYTES + _EOCD_MAX_COMMENT) :]
    position = tail.rfind(_EOCD_SIGNATURE)
    if position == -1 or len(tail) - position < _EOCD_MIN_BYTES:
        return False
    total_entries = struct.unpack_from("<H", tail, position + 10)[0]
    directory_size = struct.unpack_from("<I", tail, position + 12)[0]
    return total_entries > MAX_ZIP_MEMBERS or directory_size > MAX_CENTRAL_DIR_BYTES


def _validate_ooxml_package(payload: bytes) -> ExtractionResult | None:
    """Pre-parse zip validation; None when the package may proceed to extraction.

    Five gates, all BEFORE any document XML is materialized: (0) the RAW EOCD record's declared
    entry count / directory size (read from tail bytes — fires BEFORE ZipFile construction,
    whose directory parse is itself the resource burn) → `corrupt` 'zip-member-bound'; (1) the
    payload opens as a zip at all (non-zip garbage / broken zips → `corrupt`, exception class
    name only); (2) the summed DECLARED decompressed size of every member stays under
    MAX_DECOMPRESSED_BYTES (the raw-bytes bound — python-docx materializes every part's blob) →
    `corrupt` 'zip-expansion-bound'; (3) the POST-construction member count under
    MAX_ZIP_MEMBERS (belt-and-braces truth-after-parse vs a lying EOCD) → `corrupt`
    'zip-member-bound'; (4) the members that will be DOM-PARSED stay under MAX_XML_PARTS_BYTES
    (the parse-memory bound — lxml trees cost ~15x their XML bytes) → `corrupt`
    'xml-expansion-bound'. NO main-part-name gate: OPC resolves the main document via
    _rels/.rels, so python-docx must get its chance on any structurally sound package.
    """
    if _zip_directory_over_bound(payload):
        return ExtractionResult(None, STATUS_CORRUPT, detail="zip-member-bound")
    try:
        with zipfile.ZipFile(BytesIO(payload)) as archive:
            declared_expansion = sum(member.file_size for member in archive.infolist())
            if declared_expansion > MAX_DECOMPRESSED_BYTES:
                return ExtractionResult(None, STATUS_CORRUPT, detail="zip-expansion-bound")
            if len(archive.infolist()) > MAX_ZIP_MEMBERS:
                return ExtractionResult(None, STATUS_CORRUPT, detail="zip-member-bound")
            if _xml_parts_over_bound(archive):
                return ExtractionResult(None, STATUS_CORRUPT, detail="xml-expansion-bound")
    except Exception as zip_error:  # BadZipFile and friends — class name only, never str(exc)
        return ExtractionResult(
            None, STATUS_CORRUPT, detail=f"zipfile:{type(zip_error).__name__}"
        )
    return None


def _xml_parts_over_bound(archive: zipfile.ZipFile) -> bool:
    """True when the members that will be DOM-parsed declare more than MAX_XML_PARTS_BYTES.

    Two passes, cheapest first. Pass 1 sums the declared file_size of every *.xml / *.rels
    member straight off the zip directory — nothing is decompressed — and also rejects a
    pathological [Content_Types].xml (> MAX_CONTENT_TYPES_BYTES) before pass 2 would parse it.
    Pass 2 adds members the manifest maps to an XML content type under a NON-xml name: python-docx
    selects the part class by CONTENT TYPE, so an Override like
    '/word/media/evil.png → …styles+xml' is lxml-parsed at open despite its media suffix.
    Pass 1's bound makes pass 2's manifest parse safe.
    """
    declared_sizes = {member.filename: member.file_size for member in archive.infolist()}
    xml_named = {
        name for name in declared_sizes if name.lower().endswith(_XML_MEMBER_SUFFIXES)
    }
    xml_named_total = sum(declared_sizes[name] for name in xml_named)
    if xml_named_total > MAX_XML_PARTS_BYTES:
        return True
    if declared_sizes.get(_CONTENT_TYPES_XML, 0) > MAX_CONTENT_TYPES_BYTES:
        return True
    disguised = _xml_typed_member_names(archive) - xml_named
    disguised_total = sum(declared_sizes.get(name, 0) for name in disguised)
    return xml_named_total + disguised_total > MAX_XML_PARTS_BYTES


def _xml_typed_member_names(archive: zipfile.ZipFile) -> set[str]:
    """Member names [Content_Types].xml maps to an XML content type (Override + Default rules).

    Absent/unparseable manifest → empty set: python-docx then rejects the whole package and the
    stdlib fallback parses only word/document.xml, which the suffix pass already bounds. The
    caller has already bounded the manifest's declared size before this parse.
    """
    try:
        manifest_root = ElementTree.fromstring(archive.read(_CONTENT_TYPES_XML))
    except Exception:
        return set()
    xml_extensions = {
        (entry.get("Extension") or "").lower()
        for entry in manifest_root.iter(f"{_CONTENT_TYPES_NS}Default")
        if _is_xml_content_type(entry.get("ContentType"))
    }
    xml_typed = {
        (entry.get("PartName") or "").lstrip("/")
        for entry in manifest_root.iter(f"{_CONTENT_TYPES_NS}Override")
        if _is_xml_content_type(entry.get("ContentType"))
    }
    for member in archive.namelist():
        extension = member.rsplit(".", 1)[-1].lower() if "." in member else ""
        if extension in xml_extensions:
            xml_typed.add(member)
    return xml_typed


def _is_xml_content_type(content_type: str | None) -> bool:
    """OPC content types python-docx parses into a DOM: '+xml' suffix or a bare XML media type."""
    if not content_type:
        return False
    normalized = content_type.strip().lower()
    return normalized.endswith("+xml") or normalized in ("application/xml", "text/xml")


def _read_with_python_docx(payload: bytes) -> str:
    """Primary pass: body blocks in DOCUMENT ORDER — paragraphs as lines, tables pipe-rows.

    Document.iter_inner_content() (python-docx ≥1.1, verified on 1.2.0) yields the body's
    top-level Paragraph and Table blocks in their actual XML order, so a paragraph→table→
    paragraph document round-trips in sequence and table-cell text appears exactly once (cell
    paragraphs are reached only through their Table). Headers/footers are section parts and
    never appear in the body iteration — the deliberate boilerplate skip. Returns '' when no
    block carries any non-whitespace content.
    """
    document = Document(BytesIO(payload))
    chunks: list[str] = []
    for block in document.iter_inner_content():
        if isinstance(block, Table):
            rendered = serialize_table(_table_rows_once(block))
        else:
            rendered = block.text
        if rendered.strip():
            chunks.append(rendered)
    return "\n".join(chunks)


def _table_rows_once(table: Table) -> list[list[str]]:
    """A table's cell texts with every MERGED cell emitted exactly once.

    row.cells returns the merge-ORIGIN cell repeated for every grid slot it spans (horizontally
    per row, vertically via vMerge across rows) - naive serialization repeats 'Summary' once per
    spanned column/row (2026-06-12 review). Tracked by the identity of the underlying `w:tc`
    element (cell._tc - python-docx's stable handle for the merged region), the origin slot
    keeps the text and every other spanned slot serializes as ''.
    """
    # The ELEMENTS themselves are held in the set — not their id()s: lxml node proxies are
    # transient, so a bare id() outlives its proxy and a GC-reused address falsely dedupes an
    # UNRELATED cell (live-caught: 'Bolts' blanked). Holding the element keeps its proxy alive,
    # and lxml then returns the SAME object for the same node (identity hashing is correct).
    seen_tc_elements: set[object] = set()
    rows: list[list[str]] = []
    for row in table.rows:
        cells: list[str] = []
        for cell in row.cells:
            tc_element = cell._tc
            cells.append("" if tc_element in seen_tc_elements else cell.text)
            seen_tc_elements.add(tc_element)
        rows.append(cells)
    return rows


def _zip_xml_fallback(payload: bytes, docx_error: Exception) -> ExtractionResult:
    """python-docx crashed on a real zip — harvest word/document.xml directly (design §2.4).

    The emergency path for OOXML zips python-docx rejects (e.g. a package missing
    [Content_Types].xml): stdlib zipfile + ElementTree, one line per w:p (its descendant w:t
    texts concatenated — table-cell paragraphs included exactly once, in document order; no
    table serialization on this degraded path). Both engines failing → `corrupt` with the two
    exception class names; parsed-but-textless → `empty`.
    """
    try:
        with zipfile.ZipFile(BytesIO(payload)) as archive:
            document_xml = archive.read(_DOCUMENT_XML)
        harvested = _harvest_wordml_text(document_xml)
    except Exception as fallback_error:
        return ExtractionResult(
            None,
            STATUS_CORRUPT,
            detail=(
                f"python-docx:{type(docx_error).__name__} "
                f"xml-fallback:{type(fallback_error).__name__}"
            ),
        )
    text = sanitize_body_text(harvested)
    if not text:
        return ExtractionResult(
            None,
            STATUS_EMPTY,
            detail=f"python-docx:{type(docx_error).__name__} (xml fallback parsed, no text)",
            extractor_name=_XML_FALLBACK_NAME,
            extractor_version=_XML_FALLBACK_VERSION,
        )
    return _text_result(
        text,
        _XML_FALLBACK_NAME,
        _XML_FALLBACK_VERSION,
        detail=f"python-docx:{type(docx_error).__name__} (xml fallback)",
    )


def _harvest_wordml_text(document_xml: bytes) -> str:
    """Minimal WordprocessingML text harvest: each w:p (anywhere — body or table cell) becomes
    one line of its concatenated descendant w:t texts; empty paragraphs are dropped."""
    root = ElementTree.fromstring(document_xml)
    paragraphs = (
        "".join(node.text for node in paragraph.iter(f"{_WORDML}t") if node.text)
        for paragraph in root.iter(f"{_WORDML}p")
    )
    return "\n".join(line for line in paragraphs if line.strip())


def _text_result(
    text: str, extractor_name: str, extractor_version: str, detail: str | None = None
) -> ExtractionResult:
    """Wrap sanitized text as `extracted` — or `truncated` when it exceeds MAX_EXTRACTED_CHARS
    (the capped text IS stored; the cap detail outranks any fallback-provenance detail)."""
    if len(text) <= MAX_EXTRACTED_CHARS:
        return ExtractionResult(
            text,
            STATUS_EXTRACTED,
            detail=detail,
            extractor_name=extractor_name,
            extractor_version=extractor_version,
        )
    return ExtractionResult(
        text[:MAX_EXTRACTED_CHARS],
        STATUS_TRUNCATED,
        detail=f"capped from {len(text)} chars",
        extractor_name=extractor_name,
        extractor_version=extractor_version,
    )


