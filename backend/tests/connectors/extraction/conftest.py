"""
Role: In-memory fixture builders for the extractor tests — a classic ~600-byte single-text-object
      PDF with a CORRECT xref table (computed offsets), an image-XObject page for
      scanned-detection fixtures, a pypdf-based encryptor, a python-docx document builder
      (ordered paragraph/table blocks), an openpyxl workbook builder (typed cells, number formats,
      multi-sheet, optional injected formula cached values), and a HAND-ASSEMBLED minimal TNEF
      stream builder (signature + key + attribute records per tnefparse's wire layout: level u8,
      name u16, type u16, length u32, data, checksum u16). Self-contained: no fixture files on disk.
Used by: tests/connectors/extraction/test_pdf.py + test_docx.py + test_xlsx.py + test_tnef.py, the
         seam test (test_attachment_extractor) and the ingest integration test
         (test_email_ingest_service) import build_pdf/TEXT_PAGE_STREAM/build_docx/build_xlsx/
         build_tnef for real parseable payloads.
Depends on: pypdf (the encryptor only), python-docx (the docx builder only), openpyxl (the workbook
            builder only), compressed_rtf (the TNEF rtf-body compressor only); stdlib.
Key invariants:
  - build_pdf output is a VALID PDF (correct xref byte offsets) — pdfplumber and pypdf both parse
    it strictly, so tests never depend on parser error-recovery.
  - The image page's 1×1 XObject is drawn scaled to the FULL MediaBox (612×792) — it satisfies
    the §3.1 page-covering test by geometry, not by image payload size.
  - build_docx writes blocks in the GIVEN order (paragraph/table interleaving preserved) — the
    extractor's document-order contract is testable against it.
  - build_xlsx writes a real openpyxl workbook: each sheet's rows are appended in order, typed
    Python values ride as their native cell type (int/float/bool/datetime/str), per-cell number
    formats apply, and an optional formula spec injects a CACHED <v> into the sheet XML so the
    extractor's data_only read returns a value (openpyxl never computes formulas itself).
  - build_tnef output is a VALID TNEF stream tnefparse parses strictly (correct checksums, valid
    version attribute) — verified live against tnefparse 1.4.0; rtf bodies ride as
    LZFu-compressed MAPI_RTF_COMPRESSED properties exactly like Outlook writes them.
"""

from __future__ import annotations

import struct
import zipfile
from io import BytesIO

import docx
import openpyxl
from compressed_rtf import compress as compress_rtf
from pypdf import PdfReader, PdfWriter

# One text object, Helvetica 24pt — the classic minimal text page. Override the shown string by
# building a stream via text_page_stream().
TEXT_PAGE_STREAM = b"BT /F1 24 Tf 72 720 Td (Hello World) Tj ET"
# The 1x1 image XObject drawn over the WHOLE 612x792 page (the §3.1 page-covering geometry).
IMAGE_PAGE_STREAM = b"q 612 0 0 792 0 0 cm /Im1 Do Q"
EMPTY_PAGE_STREAM = b""


def text_page_stream(text: str) -> bytes:
    """A content stream painting `text` as one Tj string (escape PDF string delimiters)."""
    escaped = text.replace("\\", r"\\").replace("(", r"\(").replace(")", r"\)")
    return b"BT /F1 24 Tf 72 720 Td (" + escaped.encode("latin-1") + b") Tj ET"


def build_pdf(page_streams: list[bytes], with_image: bool = False) -> bytes:
    """Assemble a minimal valid PDF: one page per content stream, correct xref offsets.

    Args:
        page_streams: each entry becomes one page's content stream (use TEXT_PAGE_STREAM /
            IMAGE_PAGE_STREAM / EMPTY_PAGE_STREAM or text_page_stream()).
        with_image: register a 1×1 RGB image XObject as /Im1 in every page's resources
            (pages only SHOW it when their stream draws it, e.g. IMAGE_PAGE_STREAM).

    Returns:
        The full PDF bytes (~600 bytes for one text page).
    """
    n_pages = len(page_streams)
    font_number = 3 + 2 * n_pages
    image_number = font_number + 1 if with_image else None
    kids = " ".join(f"{3 + 2 * i} 0 R" for i in range(n_pages))

    bodies: list[bytes] = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        f"<< /Type /Pages /Kids [{kids}] /Count {n_pages} >>".encode(),
    ]
    for page_index, stream in enumerate(page_streams):
        content_number = 3 + 2 * page_index + 1
        resources = f"<< /Font << /F1 {font_number} 0 R >>"
        if image_number is not None:
            resources += f" /XObject << /Im1 {image_number} 0 R >>"
        resources += " >>"
        bodies.append(
            (
                f"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
                f"/Contents {content_number} 0 R /Resources {resources} >>"
            ).encode()
        )
        bodies.append(
            b"<< /Length " + str(len(stream)).encode() + b" >>\nstream\n" + stream + b"\nendstream"
        )
    bodies.append(b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>")
    if image_number is not None:
        pixel = b"\xff\x00\x00"  # one red RGB pixel
        bodies.append(
            b"<< /Type /XObject /Subtype /Image /Width 1 /Height 1 /ColorSpace /DeviceRGB "
            b"/BitsPerComponent 8 /Length "
            + str(len(pixel)).encode()
            + b" >>\nstream\n"
            + pixel
            + b"\nendstream"
        )

    document = bytearray(b"%PDF-1.4\n")
    offsets: list[int] = []
    for object_number, body in enumerate(bodies, start=1):
        offsets.append(len(document))
        document += f"{object_number} 0 obj\n".encode() + body + b"\nendobj\n"
    xref_position = len(document)
    entry_count = len(bodies) + 1
    document += f"xref\n0 {entry_count}\n".encode()
    document += b"0000000000 65535 f \n"
    for offset in offsets:
        document += f"{offset:010d} 00000 n \n".encode()
    document += (
        f"trailer\n<< /Size {entry_count} /Root 1 0 R >>\nstartxref\n{xref_position}\n%%EOF\n"
    ).encode()
    return bytes(document)


def build_docx(blocks: list[str | list[list[str]]]) -> bytes:
    """Assemble a real docx via python-docx: blocks land in the body in the GIVEN order.

    Args:
        blocks: each entry is either a str (one body paragraph) or a list of rows
            (list[list[str]] — one table, every row the same width).

    Returns:
        The full docx bytes (the default python-docx template; empty blocks → an empty body).
    """
    document = docx.Document()
    for block in blocks:
        if isinstance(block, str):
            document.add_paragraph(block)
            continue
        table = document.add_table(rows=len(block), cols=len(block[0]))
        for row, values in zip(table.rows, block, strict=True):
            for cell, value in zip(row.cells, values, strict=True):
                cell.text = value
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


# ── xlsx workbook assembly (openpyxl, design §2.5) ──

# One cell number-format override or a formula+cached-value spec, keyed by A1 reference. A plain
# string value is the number_format mask; a ('formula', expr, cached) tuple writes the formula AND
# injects its cached <v> into the sheet XML (openpyxl with data_only never computes formulas).
CellSpec = str | tuple[str, str, object]


def build_xlsx(
    sheets: list[tuple[str, list[list[object]]]],
    cell_specs: dict[str, dict[str, CellSpec]] | None = None,
) -> bytes:
    """Assemble a real xlsx via openpyxl: each sheet's rows are appended in order.

    Args:
        sheets: a list of (sheet_name, rows) — each row a list of native Python values
            (int/float/bool/datetime/date/str/None) written as their native cell type. An empty
            rows list makes an empty sheet.
        cell_specs: optional per-sheet {sheet_name: {ref: spec}} overrides — a str spec is a
            number_format mask applied to that cell; a ('formula', expr, cached) spec writes the
            formula to that cell AND injects its cached value into the sheet XML so the extractor's
            data_only read returns `cached` (a library workbook never opened by Excel otherwise has
            no cached value).

    Returns:
        The full xlsx bytes (openpyxl reads them with read_only + data_only).
    """
    cell_specs = cell_specs or {}
    workbook = openpyxl.Workbook()
    workbook.remove(workbook.active)  # drop the default sheet — we add named ones explicitly
    formula_caches: dict[str, dict[str, object]] = {}
    for sheet_index, (name, rows) in enumerate(sheets):
        worksheet = workbook.create_sheet(title=name)
        for row in rows:
            worksheet.append(row)
        for ref, spec in cell_specs.get(name, {}).items():
            if isinstance(spec, tuple):  # ('formula', expr, cached)
                _, expression, cached = spec
                worksheet[ref] = f"={expression}"
                formula_caches.setdefault(f"sheet{sheet_index + 1}.xml", {})[ref] = cached
            else:
                worksheet[ref].number_format = spec
    buffer = BytesIO()
    workbook.save(buffer)
    if not formula_caches:
        return buffer.getvalue()
    return _inject_formula_caches(buffer.getvalue(), formula_caches)


def _inject_formula_caches(payload: bytes, caches: dict[str, dict[str, object]]) -> bytes:
    """Inject cached <v> values into formula cells' sheet XML (the data_only read source).

    openpyxl writes a formula as `<c r="A2"><f>EXPR</f></c>` with no cached value; Excel would add
    `<v>RESULT</v>`. For a fixture the extractor can read under data_only, this splices that <v>
    after each named cell's </f>. Keyed by the worksheet member's basename (e.g. 'sheet1.xml').
    """
    out = BytesIO()
    with (
        zipfile.ZipFile(BytesIO(payload)) as source,
        zipfile.ZipFile(out, "w", zipfile.ZIP_DEFLATED) as target,
    ):
        for member in source.namelist():
            data = source.read(member)
            basename = member.rsplit("/", 1)[-1]
            if member.startswith("xl/worksheets/") and basename in caches:
                text = data.decode("utf-8")
                for ref, cached in caches[basename].items():
                    text = _splice_cached_value(text, ref, cached)
                data = text.encode("utf-8")
            target.writestr(member, data)
    return out.getvalue()


def _splice_cached_value(sheet_xml: str, ref: str, cached: object) -> str:
    """Add a `<v>cached</v>` right after the </f> of cell `ref` in one sheet's XML."""
    cached_value = "1" if cached is True else "0" if cached is False else str(cached)
    marker = f'<c r="{ref}"'
    cell_start = sheet_xml.index(marker)
    formula_end = sheet_xml.index("</f>", cell_start) + len("</f>")
    return sheet_xml[:formula_end] + f"<v>{cached_value}</v>" + sheet_xml[formula_end:]


# ── TNEF stream assembly (tnefparse wire layout, design §2.9) ──

# TNEF stream signature (0x223E9F78) — written little-endian, followed by the u16 key.
TNEF_SIGNATURE = 0x223E9F78
# Attribute name codes (mirror tnefparse.TNEF constants — hardcoded: the fixtures must keep
# building yesterday's wire format even if the library's constants ever move).
_ATT_TNEF_VERSION = 0x9006
_ATT_MAPI_PROPS = 0x9003
_ATT_BODY = 0x800C
_ATT_ATTACH_REND_DATA = 0x9002
_ATT_ATTACH_TITLE = 0x8010
_ATT_ATTACH_DATA = 0x800F
# MAPI property codes + the binary property type (tnefparse.properties / mapi).
_MAPI_RTF_COMPRESSED = 0x1009
_MAPI_BODY_HTML = 0x1013
_SZMAPI_BINARY = 0x0102
_VALID_TNEF_VERSION = 0x10000


def tnef_attribute(level: int, name: int, attr_type: int, data: bytes) -> bytes:
    """One TNEF attribute record: level u8 + name u16 + type u16 + length u32 + data + checksum.

    The checksum is the real sum-of-bytes & 0xFFFF tnefparse verifies (do_checksum=True default),
    so fixtures parse strictly — never through error tolerance.
    """
    header = struct.pack("<BHHI", level, name, attr_type, len(data))
    return header + data + struct.pack("<H", sum(data) & 0xFFFF)


def _tnef_mapi_binary_property(property_code: int, payload: bytes) -> bytes:
    """One MAPI binary property block (type u16 + name u16 + count u32 + length u32 + padded
    payload) as decode_mapi expects. The payload must not END with NUL bytes: tnefparse rstrips
    trailing NULs (padding removal) and would corrupt such a value."""
    padding = b"\x00" * (-len(payload) % 4)
    return (
        struct.pack("<HH", _SZMAPI_BINARY, property_code)
        + struct.pack("<I", 1)
        + struct.pack("<I", len(payload))
        + payload
        + padding
    )


def build_tnef(
    body: bytes | None = None,
    rtf_body: bytes | None = None,
    html_body: bytes | None = None,
    attachments: list[tuple[bytes, bytes]] | None = None,
    version: int = _VALID_TNEF_VERSION,
) -> bytes:
    """Assemble a minimal VALID TNEF stream (signature + key + attribute records).

    Args:
        body: optional plain-text body bytes (ATTBODY).
        rtf_body: optional RTF source bytes — LZFu-compressed via compressed_rtf and wrapped as
            the MAPI_RTF_COMPRESSED property, exactly how Outlook ships it.
        html_body: optional HTML body bytes (MAPI_BODY_HTML; must not end with NUL — tnefparse
            rstrips padding).
        attachments: optional (name_bytes, data_bytes) embedded files, each ≥ 2 data bytes
            (tnefparse's object loop skips trailing sub-minimum records).
        version: the TNEF version attribute value (pass a wrong one to provoke tnefparse's
            version warning for leak tests).

    Returns:
        The full TNEF bytes — tnefparse.TNEF(payload) parses them with checksums verified.
    """
    stream = struct.pack("<IH", TNEF_SIGNATURE, 0x1234)
    stream += tnef_attribute(1, _ATT_TNEF_VERSION, 0x6, struct.pack("<I", version))
    mapi_properties = b""
    property_count = 0
    if rtf_body is not None:
        mapi_properties += _tnef_mapi_binary_property(_MAPI_RTF_COMPRESSED, compress_rtf(rtf_body))
        property_count += 1
    if html_body is not None:
        mapi_properties += _tnef_mapi_binary_property(_MAPI_BODY_HTML, html_body)
        property_count += 1
    if property_count:
        stream += tnef_attribute(
            1, _ATT_MAPI_PROPS, 0x6, struct.pack("<I", property_count) + mapi_properties
        )
    if body is not None:
        stream += tnef_attribute(1, _ATT_BODY, 0x2, body)
    for name, data in attachments or []:
        # ATTATTACHRENDDATA opens a new attachment; title + data attach to it (level 2).
        stream += tnef_attribute(2, _ATT_ATTACH_REND_DATA, 0x6, b"\x00" * 14)
        stream += tnef_attribute(2, _ATT_ATTACH_TITLE, 0x7, name)
        stream += tnef_attribute(2, _ATT_ATTACH_DATA, 0x6, data)
    return stream


def encrypt_pdf(pdf_bytes: bytes, user_password: str, owner_password: str | None = None) -> bytes:
    """Re-write a PDF encrypted via pypdf (an empty user_password = owner-only lock)."""
    writer = PdfWriter()
    for page in PdfReader(BytesIO(pdf_bytes)).pages:
        writer.add_page(page)
    writer.encrypt(user_password=user_password, owner_password=owner_password)
    encrypted = BytesIO()
    writer.write(encrypted)
    return encrypted.getvalue()
