"""
Role: Unit tests for the docx extractor — document-order paragraph + table extraction (tables
      pipe-serialized via the shared serialize_table, cell text exactly once), the OLE
      password-protected probe, corrupt/zip-validation statuses, the zip-expansion (raw-bytes)
      bound, the XML-parts DOM-expansion bound (real tiny-paragraph bomb, media headroom, the
      [Content_Types].xml Override disguise, the pathological manifest), the stdlib XML-harvest
      fallback for packages python-docx rejects, truncation, C0 + surrogate sanitization, the
      vendor-log leak regression, and the NEVER-raises property. Fixtures are built in-memory
      with python-docx / stdlib zipfile (see conftest) — pure, no I/O.
Used by: pytest (tests/connectors/extraction).
Depends on: app.connectors.extraction.docx, .extraction_result, the conftest
            build_docx builder.
"""

from __future__ import annotations

import logging
import zipfile
from importlib import metadata
from io import BytesIO

import pytest

import app.connectors.extraction.docx as docx_module
from app.connectors.extraction.docx import extract_docx_text
from app.connectors.extraction.extraction_result import (
    EXTRACTION_STATUSES,
    STATUS_CORRUPT,
    STATUS_EMPTY,
    STATUS_ENCRYPTED,
    STATUS_EXTRACTED,
    STATUS_PENDING,
    STATUS_TRUNCATED,
    ExtractionResult,
)
from tests.connectors.extraction.conftest import build_docx

# 8-byte OLE compound-file magic — how a password-protected (ECMA-376-encrypted) docx arrives.
OLE_MAGIC = b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1"

# Minimal WordprocessingML for hand-built fallback zips (python-docx rejects packages without
# [Content_Types].xml; the stdlib harvest must still read them).
WORDML_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _document_xml(paragraphs: list[str]) -> str:
    body = "".join(f"<w:p><w:r><w:t>{text}</w:t></w:r></w:p>" for text in paragraphs)
    return (
        f'<?xml version="1.0"?><w:document xmlns:w="{WORDML_NS}">'
        f"<w:body>{body}</w:body></w:document>"
    )


def _zip_with(members: dict[str, str | bytes], compress: bool = False) -> bytes:
    buffer = BytesIO()
    compression = zipfile.ZIP_DEFLATED if compress else zipfile.ZIP_STORED
    with zipfile.ZipFile(buffer, "w", compression) as archive:
        for name, content in members.items():
            archive.writestr(name, content)
    return buffer.getvalue()


def _rezip_docx_with(
    payload: bytes, extra: dict[str, bytes], replace: dict[str, bytes] | None = None
) -> bytes:
    """Re-pack a real docx with members added (`extra`) and/or substituted (`replace`)."""
    replace = replace or {}
    buffer = BytesIO()
    with (
        zipfile.ZipFile(BytesIO(payload)) as source,
        zipfile.ZipFile(buffer, "w", zipfile.ZIP_DEFLATED) as target,
    ):
        for member in source.namelist():
            target.writestr(member, replace.get(member, source.read(member)))
        for name, content in extra.items():
            target.writestr(name, content)
    return buffer.getvalue()


# — text extraction (document order, tables via the shared serializer) —


def test_extract_docx_text_paragraphs_and_table_returns_extracted() -> None:
    payload = build_docx(["Intro paragraph", [["Item", "Qty"], ["Bolts", "12"]]])

    result = extract_docx_text(payload)

    assert result.status == STATUS_EXTRACTED
    assert result.text == "Intro paragraph\nItem | Qty\nBolts | 12"


def test_extract_docx_text_paragraph_table_paragraph_order_preserved() -> None:
    # iter_inner_content() yields body blocks in XML order — a table BETWEEN two paragraphs must
    # land between them, never appended after all paragraphs (the document.paragraphs +
    # document.tables two-pass trap).
    payload = build_docx(["Before the table", [["Cell A", "Cell B"]], "After the table"])

    result = extract_docx_text(payload)

    assert result.status == STATUS_EXTRACTED
    assert result.text == "Before the table\nCell A | Cell B\nAfter the table"


def test_extract_docx_text_table_cell_text_appears_exactly_once() -> None:
    # Cell paragraphs are reached ONLY through their Table block — a body traversal that also
    # walked document.paragraphs would double every cell.
    payload = build_docx([[["UniqueCellMarker", "Other"]]])

    result = extract_docx_text(payload)

    assert result.text is not None
    assert result.text.count("UniqueCellMarker") == 1


def test_extract_docx_text_success_records_python_docx_provenance() -> None:
    result = extract_docx_text(build_docx(["any text"]))

    assert result.extractor_name == "python-docx"
    assert result.extractor_version == metadata.version("python-docx")


# — empty —


def test_extract_docx_text_document_with_no_text_returns_empty_with_null_text() -> None:
    payload = build_docx([])

    result = extract_docx_text(payload)

    assert result.status == STATUS_EMPTY
    assert result.text is None  # honest NULL, never ''


def test_extract_docx_text_whitespace_only_paragraphs_return_empty() -> None:
    payload = build_docx(["   ", "\t"])

    result = extract_docx_text(payload)

    assert result.status == STATUS_EMPTY
    assert result.text is None


# — encrypted (design §2.3 policy applied to OOXML: ECMA-376 encryption = OLE wrapper) —


def test_extract_docx_text_ole_magic_returns_encrypted_with_null_text() -> None:
    payload = OLE_MAGIC + b"\x00" * 64

    result = extract_docx_text(payload)

    assert result.status == STATUS_ENCRYPTED
    assert result.text is None
    assert result.detail == "ole-container (password-protected docx or mislabeled legacy .doc)"


# — corrupt (zip validation; class names / fixed phrases only in detail) —


def test_extract_docx_text_garbage_bytes_returns_corrupt_with_class_name_detail() -> None:
    marker = "TOP-SECRET-SALARY-TABLE"

    result = extract_docx_text(f"not a zip at all {marker}".encode())

    assert result.status == STATUS_CORRUPT
    assert result.text is None
    assert result.detail == "zipfile:BadZipFile"
    assert marker not in (result.detail or "")  # payload content never echoes into detail


def test_extract_docx_text_zip_without_document_xml_returns_corrupt() -> None:
    # No presence GATE anymore (OPC resolves the main part via rels — 2026-06-12 review): a
    # plain zip degrades through BOTH engines (python-docx rejects the package, the fallback
    # finds no word/document.xml) to `corrupt` with the two class names.
    payload = _zip_with({"hello.txt": "just a plain zip"})

    result = extract_docx_text(payload)

    assert result.status == STATUS_CORRUPT
    assert result.text is None
    assert result.detail is not None
    assert result.detail.startswith("python-docx:")
    assert "xml-fallback:" in result.detail


def test_extract_docx_text_nonstandard_main_part_name_still_extracts() -> None:
    # OPC resolves the main part via _rels/.rels — Word's 'document2.xml' save quirk and
    # non-Word producers are legitimate packages python-docx parses fine; a literal-name gate
    # falsely stamped them `corrupt` (-> NULL text, bytes discarded on the live path:
    # permanent loss of a valid document; 2026-06-12 review).
    source = build_docx(["hello quirk"])
    renames = {
        "word/document.xml": "word/document2.xml",
        "word/_rels/document.xml.rels": "word/_rels/document2.xml.rels",
    }
    out = BytesIO()
    with zipfile.ZipFile(BytesIO(source)) as src, zipfile.ZipFile(out, "w") as dst:
        for member in src.namelist():
            data = src.read(member)
            if member in ("[Content_Types].xml", "_rels/.rels"):
                data = data.replace(b"word/document.xml", b"word/document2.xml")
            dst.writestr(renames.get(member, member), data)

    result = extract_docx_text(out.getvalue())

    assert result.status == STATUS_EXTRACTED
    assert result.text is not None and "hello quirk" in result.text


def test_extract_docx_text_merged_table_cells_serialized_once() -> None:
    # row.cells repeats the merge-origin cell for every spanned grid slot (and vMerge repeats
    # it per spanned row) — 'Summary' must appear exactly once, not once per slot.
    from docx import Document as DocxDocument

    document = DocxDocument()
    table = document.add_table(rows=2, cols=3)
    merged = table.cell(0, 0).merge(table.cell(0, 2))
    merged.text = "Summary"
    for column, value in enumerate(("a", "b", "c")):
        table.cell(1, column).text = value
    buffer = BytesIO()
    document.save(buffer)

    result = extract_docx_text(buffer.getvalue())

    assert result.status == STATUS_EXTRACTED
    assert result.text is not None
    assert result.text == "Summary |  | \na | b | c"  # origin slot once; spanned slots ''



@pytest.mark.parametrize(
    # Explicit ids: the raw-bytes params (a half-docx is ~18 KB) would otherwise explode the
    # auto-generated test id past Windows' 32767-char env-var limit via PYTEST_CURRENT_TEST.
    "payload",
    [
        pytest.param(b"", id="empty-bytes"),
        pytest.param(b"\x00" * 64, id="nul-run"),
        pytest.param(b"PK\x03\x04 truncated local header", id="bare-zip-magic"),
        pytest.param(
            build_docx(["mutate me"])[: len(build_docx(["mutate me"])) // 2],
            id="truncated-docx",
        ),
        pytest.param(
            _zip_with({"word/document.xml": "<w:document broken"}), id="broken-inner-xml"
        ),
    ],
)
def test_extract_docx_text_mutated_payloads_never_raise(payload: bytes) -> None:
    result = extract_docx_text(payload)

    assert isinstance(result, ExtractionResult)
    assert result.status in EXTRACTION_STATUSES and result.status != STATUS_PENDING


# — the stdlib XML-harvest fallback (design §2.4's emergency path) —


def test_extract_docx_text_package_without_content_types_falls_back_to_xml_harvest() -> None:
    # python-docx refuses a zip missing [Content_Types].xml (KeyError) — but word/document.xml
    # is right there; the stdlib harvest must recover its text rather than calling it corrupt.
    payload = _zip_with(
        {"word/document.xml": _document_xml(["Recovered first line", "Recovered second line"])}
    )

    result = extract_docx_text(payload)

    assert result.status == STATUS_EXTRACTED
    assert result.text == "Recovered first line\nRecovered second line"
    assert result.extractor_name == "docx-zip-xml"
    assert result.detail is not None and "python-docx:KeyError" in result.detail


def test_extract_docx_text_fallback_textless_document_returns_empty() -> None:
    payload = _zip_with({"word/document.xml": _document_xml([])})

    result = extract_docx_text(payload)

    assert result.status == STATUS_EMPTY
    assert result.text is None


def test_extract_docx_text_fallback_unparseable_xml_returns_corrupt_with_both_class_names() -> (
    None
):
    payload = _zip_with({"word/document.xml": "<w:document never closed"})

    result = extract_docx_text(payload)

    assert result.status == STATUS_CORRUPT
    assert result.detail is not None
    assert "python-docx:" in result.detail and "xml-fallback:" in result.detail


# — truncation (MAX_EXTRACTED_CHARS shared with the PDF extractor; capped text IS stored) —


def test_extract_docx_text_above_cap_truncates_and_stores_capped_text(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    monkeypatch.setattr(docx_module, "MAX_EXTRACTED_CHARS", 12)
    payload = build_docx(["Hello World of documents"])  # 24 chars

    result = extract_docx_text(payload)

    assert result.status == STATUS_TRUNCATED
    assert result.text == "Hello World "  # capped text IS stored
    assert result.detail == "capped from 24 chars"


# — sanitization (sanitize_body_text: same rules as email bodies — storable as UTF-8, period) —


def test_extract_docx_text_strips_c0_and_nul_from_engine_output(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    # lxml refuses to WRITE control chars, so a fixture docx cannot carry them — stub the engine
    # boundary like the PDF tests do: WHATEVER python-docx yields must pass the email-body C0/NUL
    # rules before storage.
    monkeypatch.setattr(
        docx_module, "_read_with_python_docx", lambda payload: "Hel\x07lo\x00 World\r\nnext"
    )

    result = extract_docx_text(build_docx(["replaced by the stub"]))

    assert result.status == STATUS_EXTRACTED
    assert result.text == "Hello World\nnext"  # BEL + NUL stripped, CRLF → LF


def test_extract_docx_text_strips_lone_surrogates_so_text_is_utf8_storable(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    monkeypatch.setattr(
        docx_module, "_read_with_python_docx", lambda payload: "Invoice \ud83d total \udfff due"
    )

    result = extract_docx_text(build_docx(["replaced by the stub"]))

    assert result.status == STATUS_EXTRACTED
    assert result.text == "Invoice  total  due"  # surrogate-free
    result.text.encode("utf-8")  # must not raise — the storability guarantee asyncpg relies on


# — vendor-log payload leak (same regression class as pypdf's header echo, pdf.py) —


def test_extract_docx_text_corrupt_payload_bytes_never_reach_logs(
    caplog: pytest.LogCaptureFixture,
) -> None:
    # The secret must actually REACH the parsers (2026-06-12 review: a raw-bytes payload dies
    # at the zip gate and never exercises python-docx/lxml/ElementTree) — so: a structurally
    # valid package whose word/document.xml is BROKEN XML carrying the marker. Both engines
    # parse-and-crash on those bytes; no log record — including formatted tracebacks
    # (caplog.text covers exc_text) — may contain payload content.
    marker = "SECRET-SALARY-TABLE-DO-NOT-DISTRIBUTE"
    source = build_docx(["placeholder"])
    out = BytesIO()
    with zipfile.ZipFile(BytesIO(source)) as src, zipfile.ZipFile(out, "w") as dst:
        for member in src.namelist():
            if member == "word/document.xml":
                dst.writestr(member, f"<broken {marker} <<not-xml".encode())
            else:
                dst.writestr(member, src.read(member))

    with caplog.at_level(logging.DEBUG):
        result = extract_docx_text(out.getvalue())

    assert result.status == STATUS_CORRUPT
    assert result.text is None
    assert marker not in (result.detail or "")
    assert marker not in caplog.text  # full formatted records, tracebacks included
