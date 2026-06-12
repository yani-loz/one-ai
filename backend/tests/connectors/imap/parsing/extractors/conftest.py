"""
Role: In-memory fixture builders for the extractor tests — a classic ~600-byte single-text-object
      PDF with a CORRECT xref table (computed offsets), an image-XObject page for
      scanned-detection fixtures, a pypdf-based encryptor, and a python-docx document builder
      (ordered paragraph/table blocks). Self-contained: no fixture files on disk, no network.
Used by: tests/connectors/imap/parsing/extractors/test_pdf.py + test_docx.py, the seam test
         (test_attachment_extractor) and the ingest integration test (test_email_ingest_service)
         import build_pdf/TEXT_PAGE_STREAM/build_docx for real parseable payloads.
Depends on: pypdf (the encryptor only), python-docx (the docx builder only); stdlib.
Key invariants:
  - build_pdf output is a VALID PDF (correct xref byte offsets) — pdfplumber and pypdf both parse
    it strictly, so tests never depend on parser error-recovery.
  - The image page's 1×1 XObject is drawn scaled to the FULL MediaBox (612×792) — it satisfies
    the §3.1 page-covering test by geometry, not by image payload size.
  - build_docx writes blocks in the GIVEN order (paragraph/table interleaving preserved) — the
    extractor's document-order contract is testable against it.
"""

from __future__ import annotations

from io import BytesIO

import docx
from pypdf import PdfReader, PdfWriter

# One text object, Helvetica 24pt — the classic minimal text page. Override the shown string by
# building a stream via text_page_stream().
TEXT_PAGE_STREAM = b"BT /F1 24 Tf 72 720 Td (Hello World) Tj ET"
# The 1x1 image XObject drawn over the WHOLE 612x792 page (the §3.1 page-covering geometry).
IMAGE_PAGE_STREAM = b"q 612 0 0 792 0 0 cm /Im1 Do Q"
EMPTY_PAGE_STREAM = b""


def text_page_stream(text: str) -> bytes:
    """A content stream painting `text` as one Tj string (escape PDF string delimiters)."""
    escaped = text.replace("\\", r"\\").replace("(", r"\(").replace(")", r"\)")
    return b"BT /F1 24 Tf 72 720 Td (" + escaped.encode("latin-1") + b") Tj ET"


def build_pdf(page_streams: list[bytes], with_image: bool = False) -> bytes:
    """Assemble a minimal valid PDF: one page per content stream, correct xref offsets.

    Args:
        page_streams: each entry becomes one page's content stream (use TEXT_PAGE_STREAM /
            IMAGE_PAGE_STREAM / EMPTY_PAGE_STREAM or text_page_stream()).
        with_image: register a 1×1 RGB image XObject as /Im1 in every page's resources
            (pages only SHOW it when their stream draws it, e.g. IMAGE_PAGE_STREAM).

    Returns:
        The full PDF bytes (~600 bytes for one text page).
    """
    n_pages = len(page_streams)
    font_number = 3 + 2 * n_pages
    image_number = font_number + 1 if with_image else None
    kids = " ".join(f"{3 + 2 * i} 0 R" for i in range(n_pages))

    bodies: list[bytes] = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        f"<< /Type /Pages /Kids [{kids}] /Count {n_pages} >>".encode(),
    ]
    for page_index, stream in enumerate(page_streams):
        content_number = 3 + 2 * page_index + 1
        resources = f"<< /Font << /F1 {font_number} 0 R >>"
        if image_number is not None:
            resources += f" /XObject << /Im1 {image_number} 0 R >>"
        resources += " >>"
        bodies.append(
            (
                f"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
                f"/Contents {content_number} 0 R /Resources {resources} >>"
            ).encode()
        )
        bodies.append(
            b"<< /Length " + str(len(stream)).encode() + b" >>\nstream\n" + stream + b"\nendstream"
        )
    bodies.append(b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>")
    if image_number is not None:
        pixel = b"\xff\x00\x00"  # one red RGB pixel
        bodies.append(
            b"<< /Type /XObject /Subtype /Image /Width 1 /Height 1 /ColorSpace /DeviceRGB "
            b"/BitsPerComponent 8 /Length "
            + str(len(pixel)).encode()
            + b" >>\nstream\n"
            + pixel
            + b"\nendstream"
        )

    document = bytearray(b"%PDF-1.4\n")
    offsets: list[int] = []
    for object_number, body in enumerate(bodies, start=1):
        offsets.append(len(document))
        document += f"{object_number} 0 obj\n".encode() + body + b"\nendobj\n"
    xref_position = len(document)
    entry_count = len(bodies) + 1
    document += f"xref\n0 {entry_count}\n".encode()
    document += b"0000000000 65535 f \n"
    for offset in offsets:
        document += f"{offset:010d} 00000 n \n".encode()
    document += (
        f"trailer\n<< /Size {entry_count} /Root 1 0 R >>\nstartxref\n{xref_position}\n%%EOF\n"
    ).encode()
    return bytes(document)


def build_docx(blocks: list[str | list[list[str]]]) -> bytes:
    """Assemble a real docx via python-docx: blocks land in the body in the GIVEN order.

    Args:
        blocks: each entry is either a str (one body paragraph) or a list of rows
            (list[list[str]] — one table, every row the same width).

    Returns:
        The full docx bytes (the default python-docx template; empty blocks → an empty body).
    """
    document = docx.Document()
    for block in blocks:
        if isinstance(block, str):
            document.add_paragraph(block)
            continue
        table = document.add_table(rows=len(block), cols=len(block[0]))
        for row, values in zip(table.rows, block, strict=True):
            for cell, value in zip(row.cells, values, strict=True):
                cell.text = value
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def encrypt_pdf(pdf_bytes: bytes, user_password: str, owner_password: str | None = None) -> bytes:
    """Re-write a PDF encrypted via pypdf (an empty user_password = owner-only lock)."""
    writer = PdfWriter()
    for page in PdfReader(BytesIO(pdf_bytes)).pages:
        writer.add_page(page)
    writer.encrypt(user_password=user_password, owner_password=owner_password)
    encrypted = BytesIO()
    writer.write(encrypted)
    return encrypted.getvalue()
